import docx
import plotly.graph_objects as go
import helpers
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn


def set_cell_margins(cell, **kwargs):
    """
    cell:  actual cell instance you want to modify
    usage:
        set_cell_margins(cell, top=50, start=50, bottom=50, end=50)

    provided values are in twentieths of a point (1/1440 of an inch).
    read more here: http://officeopenxml.com/WPtableCellMargins.php
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for m in ["top", "start", "bottom", "end"]:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

    tcPr.append(tcMar)


class ModelData:

    def __init__(self, _question_count=74, _model_area_count=6, _model_layers_count=3, _weight_tab="Baseline",
                 _question_db="Product Role Compass v4.9.xlsx",
                 _question_path="/Users/macluky/Library/CloudStorage/OneDrive-SharedLibraries-ExpandiorAcademyB.V/Expandior Team - Documents/Product/Compass (Maturity Scan Personal)"):
        self.question_db = _question_db
        self.question_path = _question_path
        self.weight_tab = _weight_tab
        self.model_layers_count = _model_layers_count
        self.question_count = _question_count
        self.area_count = _model_area_count
        self.recommendations = 18
        self.file = self.question_path + "/" + self.question_db


class ResultData:

    def __init__(self, _nr_of_candidates=8,
                 _result_path="/Users/macluky/Library/CloudStorage/OneDrive-SharedLibraries-ExpandiorAcademyB.V/Expandior Team - Documents/Product/Compass (Maturity Scan Personal)",
                 _result_db="responses.xlsx"):
        # nr_candidates is still hardcoded, should extract from file
        self.nr_of_candidates = _nr_of_candidates
        self.result_path = _result_path
        self.result_db = _result_db
        self.file = self.result_path + "/" + self.result_db


class ReportGenerator:

    def __init__(self,
                 _file="/Users/macluky/Library/CloudStorage/OneDrive-SharedLibraries-ExpandiorAcademyB.V/Expandior Team - Documents/Product/Compass (Maturity Scan Personal)/ReportTemplate.docx",
                 _results_parser=None, _advice_parser=None, _weight_parser=None, _question_parser=None):
        self.file = _file
        self.rp = _results_parser
        self.ap = _advice_parser
        self.wp = _weight_parser
        self.qp = _question_parser
        self.doc = docx.Document(self.file)
        self.results = []
        self.name = "undefined"
        self.email = "undefined"

    # handle multiple candidates
    def select_candidate(self, i):
        self.results = self.rp.candidates[i].answers
        self.name = self.rp.candidates[i].name
        self.email = self.rp.candidates[i].email

    def count_of_candidates(self):
        return len(self.rp.candidates)

    # drawing stuff
    def create_spider(self, depth):
        labels = helpers.all_labels_at_level(self.results, depth)
        fig = go.Figure()
        scores = []
        for label in labels:
            score = helpers.calc_score_for_label(self.results, label)
            scores.append(score)
        fig.add_trace(go.Scatterpolar(
            r=scores,
            theta=labels,
            fill='toself',
            name="Results",
            fillcolor="#01FFCC",
            opacity=0.6,
            line=dict(color="limegreen")
        ))
        weights = []
        for label in labels:
            if depth > 1:
                # we don't support that level of granularity
                weight = helpers.weight_for_parent_label(self.wp, self.results, label)
            else:
                weight = self.wp.weight_label(label)
            weights.append(weight)
        fig.add_trace(go.Scatterpolar(
            r=weights,
            theta=labels,
            fill='toself',
            name="Reference",
            fillcolor="#0A1765",
            opacity=0.6,
            line=dict(color="darkblue")
        ))
        fig.update_layout(
            polar=dict(
                radialaxis=dict(
                    visible=True
                ),
            ),
            margin=dict(l=30, r=30, t=30, b=30),
            showlegend=False
        )
        # flip the order of the graphs
        fig.data = (fig.data[1], fig.data[0])
        return fig

    def create_bar_graph_for_labels(self, labels):
        fig = go.Figure()
        scores = []
        for label in labels:
            score = helpers.calc_score_for_label(self.results, label)
            scores.append(score)
        fig.add_bar(y=scores, x=labels, name="Your score", marker_color="#01FFCC")

        weights = []
        for label in labels:
            weight = self.wp.weight_label(label)
            weights.append(weight)
        fig.add_bar(y=weights, x=labels, name="Reference", marker_color="#0A1765")

        # flip the order of the graphs
        fig.data = (fig.data[1], fig.data[0])
        fig.update_yaxes(title_text="Average score")
        fig.update_layout(showlegend=False, legend_title_text="Results", margin=dict(l=20, r=20, t=20, b=20))

        return fig

    def replace_tag_with_bar(self, tag, label):
        tempfile = "/Users/macluky/Downloads/temp.png"
        para = self.find_tag(tag)
        if para is None:
            print("Can't find tag: " + tag + " in document template")
        else:
            sub_labels = helpers.sub_labels_of_label(self.rp, self.results, label)
            fig = self.create_bar_graph_for_labels(sub_labels)
            fig.write_image(tempfile, format='png', width=400, height=300, scale=0.84)
            helpers.substitute_image_placeholder(para, tempfile)

    def replace_tag_with_spider_of_depth(self, tag, depth):
        tempfile = "/Users/macluky/Downloads/temp.png"
        para = self.find_tag(tag)
        if para is None:
            print("Can't find: " + tag + " in document template")
        else:
            fig = self.create_spider(depth)
            # obsolete? ended up with the same dimentions
            if depth == 2:
                fig.write_image(tempfile, format='png', width=580, height=400, scale=0.84)
            else:
                fig.write_image(tempfile, format='png', width=580, height=400, scale=0.84)

            helpers.substitute_image_placeholder(para, tempfile)
            # print("Replaced tag: " + tag + " with spider of depth "+str(depth))

    # text replacement
    def set_name(self):
        self.replace_tag_with_text("<Name>", self.name, True)
        print("Setting name to: " + self.name)

    def strip_tag(self, tag):
        para = self.find_tag(tag, False)
        if para:
            para.text = para.text.replace(tag, "")

    def replace_tag_with_text(self, tag, text, exact_match=True):
        para = self.find_tag(tag, exact_match)
        if para is None:
            print("Can't replace: " + tag + " in document template")
        else:
            # print("Replace: " + tag + " with " + text)
            para.text = text

    def find_tag(self, tag, exact_match=True):
        found = False
        for para in self.doc.paragraphs:
            if exact_match:
                if para.text == tag:
                    return para
            else:
                # if para.text.startswith(tag):
                if para.text.strip().startswith(tag):
                    return para
        if not found:
            print("Can't find: " + tag + " in document template")
        return None

    def select_tag_text_based_on_score(self, tag, score, ref):
        if score < ref:
            # print("In "+tag+" you are scoring below the reference mark " + str(score-ref_level))
            self.replace_tag_with_text("<" + tag + " Higher>", "", False)
            self.strip_tag("<" + tag + " Lower>")
        else:
            self.replace_tag_with_text("<" + tag + " Lower>", "", False)
            self.strip_tag("<" + tag + " Higher>")

    # generic
    def save(self, path):
        filename = "Compass Report for " + self.email + ".docx"
        print("Saving: [" + filename + "]")
        self.doc.save(path + "/" + filename)

    # the actual analysis
    def add_graphs_with_ref_level(self):
        # overall spider
        self.replace_tag_with_spider_of_depth("<Overall Score>", 0)

        # level 1 spider
        self.replace_tag_with_spider_of_depth("<Domain Score>", 1)

        # level 1 alternative (bar charts)
        for label in helpers.all_labels_at_level(self.results, 0):
            tag = "<" + label + ">"
            self.replace_tag_with_bar(tag, label)

        # level 2 spider
        if self.qp.levels > 2:
            self.replace_tag_with_spider_of_depth("<Details Score>", 2)

    def score_analysis(self, model_level_count):
        score = helpers.score_at_level(results=self.results, level=0)
        self.select_tag_text_based_on_score("Overall", score, self.wp.weight_label(None))

        # figure out where they can improve the most
        for label in helpers.all_labels_at_level(self.results, 1):
            score = helpers.calc_score_for_label(self.results, label)
            self.select_tag_text_based_on_score(label, score, self.wp.weight_label(label))
            if model_level_count > 2:
                tag = "<" + label + " Areas>"
                if score < self.wp.weight_label(label):
                    sub_labels = helpers.sub_labels_of_label(self.rp, self.results, label)
                    if sub_labels is not None:
                        sep = ", "
                        self.replace_tag_with_text(tag, "Area(s) of concern: " + sep.join(sub_labels))
                else:
                    self.replace_tag_with_text(tag, "")

    def add_recommendations(self):
        max_depth = self.ap.max_model_layers - 1
        total_advice = ""
        for label in helpers.all_labels_at_level(self.results, max_depth):
            score = helpers.calc_score_for_label(self.results, label)
            ref_level = helpers.weight_for_parent_label(self.wp, self.results, label)
            if score < ref_level:
                advice = self.ap.advice_for_label_at_depth_with_reference(label, max_depth, ref_level)
                if advice is not None:
                    total_advice += "To improve " + label + ": " + advice + "\n"
        self.replace_tag_with_text("<Recommendations>", total_advice)

    def add_feedback(self, responses):
        # add grid table
        q_count = len(responses[0])
        r_count = len(responses)
        table = self.doc.add_table(rows=q_count + 1, cols=r_count + 1, style="Table Grid")

        # access first row's cells
        heading_row = table.rows[0].cells

        #TODO: formatting of the table and cells

        # add headings
        heading_row[0].text = "Question"
        for i in range(0, r_count):
            response = responses[i]
            heading_row[1 + i].text = response["name"]

        questions = list(responses[0].keys())
        for i in range(0, q_count):
            cell = table.cell(row_idx=1 + i, col_idx=0)
            cell.text = questions[i]

        col = 1
        for response in responses:
            answers = list(response.values())
            for i in range(0, q_count):
                if i >= len(answers):
                    break
                cell = table.cell(row_idx=1 + i, col_idx=col)
                cell.text = str(answers[i])
            col += 1


